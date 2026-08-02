from pathlib import Path
from zipfile import ZipFile
from docx import Document
import re

ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / "documentacion" / "PROYECTO_FIREFIGHTER_UNIFICADO.docx"
ENV = ROOT / ".env.docker"

with ZipFile(DOCX) as z:
    bad = z.testzip()
    xml = z.read("word/document.xml").decode("utf-8", "ignore")
    media = [n for n in z.namelist() if n.startswith("word/media/")]

doc = Document(DOCX)
paragraph_text = "\n".join(p.text for p in doc.paragraphs)

secret_hits = 0
sensitive_keys = {
    "POSTGRES_PASSWORD",
    "IOT_NODE_KEY",
    "JWT_SECRET_KEY",
    "MAIL_PASSWORD",
    "PAYPAL_CLIENT_ID",
    "PAYPAL_CLIENT_SECRET",
    "POCKETBASE_TOKEN",
    "NGROK_AUTHTOKEN",
}
if ENV.exists():
    for line in ENV.read_text(encoding="utf-8", errors="ignore").splitlines():
        if "=" not in line or line.lstrip().startswith("#"):
            continue
        key, value = line.split("=", 1)
        if key.strip() not in sensitive_keys:
            continue
        value = value.strip()
        if len(value) >= 12 and value in xml:
            secret_hits += 1

checks = {
    "zip_integrity": bad is None,
    "paragraphs": len(doc.paragraphs),
    "tables": len(doc.tables),
    "inline_shapes": len(doc.inline_shapes),
    "embedded_media": len(media),
    "heading1": sum(1 for p in doc.paragraphs if p.style.name == "Heading 1"),
    "figure_mentions": len(re.findall(r"Figura\s+\d+", paragraph_text)),
    "algorithm_mentions": len(re.findall(r"Algoritmo\s+\d+", paragraph_text)),
    "secret_values_found": secret_hits,
    "has_architecture": "Arquitectura final del sistema" in paragraph_text,
    "has_iot": "Integración IoT ESP32 post-incendio" in paragraph_text,
    "has_compliance_matrix": "Matriz de cumplimiento del enunciado" in paragraph_text,
}

for key, value in checks.items():
    print(f"{key}: {value}")

if bad is not None or secret_hits:
    raise SystemExit(1)
