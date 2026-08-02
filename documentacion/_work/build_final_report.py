from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
BASE = ROOT / "documentacion" / "PROYECTO_FIREFIGHTER_BASE.docx"
OUT = ROOT / "documentacion" / "PROYECTO_FIREFIGHTER_UNIFICADO.docx"
FIG = ROOT / "documentacion" / "figuras"

NAVY = "102A43"
BLUE = "2F6BFF"
RED = "D9413A"
GREEN = "18875D"
GOLD = "D99000"
LIGHT = "EAF1F7"
PALE = "F7F9FB"
WHITE = "FFFFFF"
GRAY = "52677D"

shutil.copy2(BASE, OUT)
doc = Document(OUT)

# Keep the institutional cover and the original automatic TOC, but replace the
# historical second-partial body with one coherent final report.
first_body = next(
    p for p in doc.paragraphs
    if p.text.strip().upper() == "ÍNDICE"
)
body = doc._element.body
remove_mode = False
for child in list(body):
    if child is first_body._p:
        remove_mode = True
    if remove_mode and child.tag != qn("w:sectPr"):
        body.remove(child)

def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def set_cell_text(cell, text, bold=False, color=NAVY, size=9, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    margins(cell)

def add_heading(text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p

def add_body(text, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p

def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Paragraph")
        p.style = doc.styles["List Bullet"] if "List Bullet" in [s.name for s in doc.styles] else doc.styles["List Paragraph"]
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        r.font.size = Pt(11)

def add_table(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        set_cell_text(hdr.cells[i], h, True, WHITE, 9, WD_ALIGN_PARAGRAPH.CENTER)
        shade(hdr.cells[i], NAVY)
        if widths:
            hdr.cells[i].width = Inches(widths[i])
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, False, NAVY, 8.5)
            shade(cells[i], WHITE if ri % 2 == 0 else LIGHT)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table

def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "Actualizar campo en Word"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, sep, txt, end])

def add_figure(filename, title, explanation, width=6.35, reset=None):
    p_intro = add_body(explanation)
    p_intro.paragraph_format.keep_with_next = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    r = p.add_run()
    r.add_picture(str(FIG / filename), width=Inches(width))
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_with_next = False
    rr = cap.add_run("Figura ")
    rr.bold = True
    instr = "SEQ Figura \\* ARABIC"
    if reset is not None:
        instr += f" \\r {reset}"
    add_field(cap, instr)
    r2 = cap.add_run(f". {title}")
    r2.italic = True
    return p

def add_algorithm(number, title, code, explanation):
    add_body(explanation)
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    a = cap.add_run(f"Algoritmo {number}. ")
    a.bold = True
    b = cap.add_run(title)
    b.italic = True
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0,0)
    shade(cell, "F1F4F7")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for idx, line in enumerate(code.splitlines()):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(NAVY)
    margins(cell, 120, 140, 120, 140)
    doc.add_paragraph()

def add_status_box(label, text, color):
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width = Inches(1.25)
    t.columns[1].width = Inches(5.05)
    set_cell_text(t.cell(0,0), label, True, WHITE, 9, WD_ALIGN_PARAGRAPH.CENTER)
    shade(t.cell(0,0), color)
    set_cell_text(t.cell(0,1), text, False, NAVY, 9)
    shade(t.cell(0,1), PALE)
    doc.add_paragraph()

# Update only the identifying fields on the original cover.
for p in doc.paragraphs[:25]:
    if "TEMA" in p.text:
        p.text = "TEMA\t: \tSistema distribuido de gestión bomberil"
    elif "FECHA DE ENTREGA" in p.text:
        p.text = "FECHA DE ENTREGA\t: \t29/07/2026"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(14)
r = p.add_run("ÍNDICE")
r.bold = True
r.font.name = "Arial"
r.font.size = Pt(14)
p = doc.add_paragraph()
add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.page_break_before = True
p.paragraph_format.space_after = Pt(14)
r = p.add_run("ÍNDICE DE FIGURAS")
r.bold = True
r.font.name = "Arial"
r.font.size = Pt(14)
p = doc.add_paragraph()
add_field(p, 'TOC \\h \\z \\c "Figura"')

p = add_heading("Resumen", 1)
p.paragraph_format.page_break_before = True
add_body(
    "El proyecto consolida una plataforma distribuida para gestión de emergencias bomberiles. "
    "El ciudadano registra incidentes con geolocalización y evidencia; el administrador recibe la alerta "
    "en tiempo real y despacha una unidad operada por personal autenticado; el operador navega al sitio, "
    "gestiona el estado y documenta el cierre. El sistema integra Angular 21.2, Spring Boot 4.0.6, "
    "PostgreSQL, STOMP sobre WebSocket, Docker Swarm, PocketBase, PayPal Sandbox, autenticación reforzada, "
    "Ollama con llama3.2 y un nodo IoT ESP32 para evaluación post-incendio."
)
add_body(
    "La solución se despliega con dos réplicas del frontend y dos del backend, mientras la base de datos "
    "permanece en una sola instancia para centralizar la consistencia. Las pruebas ejecutadas registraron "
    "12 casos exitosos en el backend y 7 en el frontend. El almacenamiento multimedia se separó de la "
    "base relacional y se añadió integridad SHA-256. Como limitación técnica, la asincronía actual emplea "
    "un ThreadPoolTaskExecutor con cola en memoria; no existe todavía un broker persistente externo."
)
add_body(
    "Palabras clave: aplicaciones distribuidas, Docker Swarm, WebSocket, ESP32, PocketBase, PayPal Sandbox, "
    "TOTP, Ollama, PostgreSQL."
)

add_heading("Introducción", 1)
add_body(
    "La gestión de una emergencia requiere que información, recursos y decisiones se coordinen con la menor "
    "latencia posible. Este proyecto integra en una sola solución el reporte ciudadano, el despacho central, "
    "la operación de unidades, el cierre documental, el monitoreo post-incendio y servicios complementarios "
    "de prevención. La arquitectura distribuye responsabilidades sin fragmentar la experiencia del usuario."
)
add_body(
    "La arquitectura responde a un caso realista. Los reportes ciudadanos son variables e "
    "impredecibles; las unidades deben asignarse una sola vez; los estados tienen que propagarse sin "
    "recargar la página; la evidencia debe conservar su integridad; y la lectura de aire post-incendio "
    "se asocia con el informe operativo. La distribución se utiliza donde aporta disponibilidad y "
    "se conserva una única fuente de verdad transaccional."
)

add_heading("Descripción del problema", 1)
add_body(
    "Los procedimientos tradicionales de atención dependen de llamadas, radio y actualización manual. Esto "
    "dificulta recibir evidencia visual y ubicación precisa, identificar qué unidades tienen realmente un "
    "operador disponible, evitar asignaciones simultáneas y conservar una bitácora verificable. Después de un "
    "incendio también existe incertidumbre sobre la calidad del aire, cuya observación suele quedar separada del reporte."
)
add_body(
    "Desde el punto de vista informático, una aplicación centralizada en un solo proceso se convierte en un punto "
    "único de fallo. Al mismo tiempo, replicar servicios sin coordinar estado, WebSocket, archivos y transacciones "
    "puede generar inconsistencias. El problema consiste en diseñar una plataforma que distribuya la carga y se "
    "recupere de la caída de tareas, pero mantenga una sola verdad operacional."
)

add_heading("Justificación", 1)
add_body(
    "La plataforma aporta valor social porque reduce pasos entre el ciudadano y la central, mejora la visibilidad "
    "del administrador y guía al operador. Académicamente, permite aplicar REST, WebSocket, concurrencia, JWT, "
    "OAuth 2.0, transacciones, réplicas, almacenamiento externo, pagos Sandbox, autenticación multifactor, IA local "
    "e IoT dentro de un caso cohesionado de Aplicaciones Distribuidas."
)

add_heading("Objetivos", 1)
add_heading("Objetivo general", 2)
add_body(
    "Ampliar y validar la aplicación distribuida de gestión bomberil mediante clúster, comunicación "
    "síncrona y en tiempo real, seguridad multifactor, almacenamiento privado, pagos de prueba, inteligencia "
    "artificial local y telemetría IoT, conservando la consistencia operativa del sistema."
)
add_heading("Objetivos específicos", 2)
add_bullets([
    "Desplegar frontend y backend como servicios replicados de Docker Swarm y verificar su estado deseado.",
    "Integrar PayPal Sandbox para el plan Premium de prevención, validando monto, moneda, orden y captura.",
    "Implementar OTP por correo y TOTP compatible con aplicaciones autenticadoras.",
    "Guardar evidencias en PocketBase con metadatos relacionales e integridad SHA-256.",
    "Vincular la telemetría del ESP32 al reporte del operador y a la bitácora final.",
    "Incorporar un copiloto local Ollama limitado a información operativa autorizada.",
    "Ejecutar pruebas automatizadas y documentar de forma transparente las brechas pendientes."
])

add_heading("Arquitectura final del sistema", 1)
add_figure(
    "arquitectura-final.png",
    "Arquitectura distribuida final de la plataforma.",
    "La Figura 1 presenta la topología consolidada. El routing mesh de Swarm expone los servicios replicados, "
    "Spring Boot concentra las reglas de negocio y PostgreSQL actúa como fuente transaccional. PocketBase, "
    "Ollama, PayPal, correo y ESP32 se integran mediante interfaces controladas.",
    6.35, 1
)
add_body(
    "La decisión de mantener una instancia de PostgreSQL responde al enunciado y evita divergencias entre "
    "réplicas. El backend no conserva sesión HTTP local: autentica mediante JWT y persiste el estado compartido. "
    "No obstante, el broker STOMP configurado es simple y vive dentro de cada réplica. Para compensar esta "
    "limitación en el prototipo, el cliente reconcilia periódicamente el estado por REST; en producción se "
    "recomienda un relay STOMP externo."
)
add_table(
    ["Capa", "Componente", "Instancias", "Responsabilidad"],
    [
        ["Presentación", "Angular 21.2 + Nginx", "2", "Interfaces de ciudadano, administrador y operador."],
        ["Negocio", "Spring Boot 4.0.6", "2", "REST, seguridad, transacciones, STOMP e integraciones."],
        ["Datos", "PostgreSQL 17", "1", "Estado transaccional y trazabilidad."],
        ["Archivos", "PocketBase", "1", "Binarios multimedia privados."],
        ["IA", "Ollama / llama3.2", "1", "Copiloto operativo local."],
        ["IoT", "ESP32 + MQ-2", "1 nodo", "Telemetría de humo/gas post-incendio."],
        ["Exposición", "ngrok", "1", "Túnel HTTPS para el dispositivo IoT."],
    ],
    [1.0, 1.8, 0.75, 2.8]
)

add_heading("Tecnologías utilizadas", 1)
add_table(
    ["Tecnología", "Versión / modalidad", "Uso en el proyecto"],
    [
        ["Angular", "21.2", "SPA, componentes, guards, formularios y cliente STOMP."],
        ["Spring Boot", "4.0.6 / Java 25", "API REST, JPA, seguridad, correo, pagos e IoT."],
        ["PostgreSQL", "17 Alpine", "Persistencia relacional ACID."],
        ["Docker Swarm", "Modo Swarm", "Orquestación, réplicas, red overlay e ingress."],
        ["PocketBase", "Contenedor local", "Archivos y registros de evidencia."],
        ["PayPal", "Orders API v2 Sandbox", "Creación, aprobación y captura de pago."],
        ["Ollama", "llama3.2 local", "Resumen y consultas operativas acotadas."],
        ["ESP32", "DOIT DevKit V1", "Adquisición analógica y envío HTTPS."],
        ["STOMP/WebSocket", "Broker simple + SockJS", "Difusión en vivo de incidentes y estados."],
    ],
    [1.35, 1.6, 3.1]
)

add_heading("Marco teórico", 1)
add_heading("Orquestación y réplicas con Docker Swarm", 2)
add_body(
    "Swarm administra servicios mediante un estado deseado. Cada réplica corresponde a una tarea independiente; "
    "si una tarea termina, el orquestador crea otra para recuperar el número configurado. El routing mesh acepta "
    "solicitudes en el puerto publicado y las conduce hacia una tarea activa. En este proyecto la disponibilidad "
    "frente a la caída de un contenedor queda demostrada, pero el clúster posee un solo nodo físico; por tanto, "
    "no tolera el apagado total del equipo anfitrión."
)
add_heading("Comunicación síncrona, tiempo real y asincronía", 2)
add_body(
    "REST implementa el patrón solicitud–respuesta para comandos y consultas que requieren confirmación inmediata. "
    "STOMP sobre WebSocket mantiene un canal bidireccional y publica eventos a suscriptores. El envío de correo se "
    "ejecuta mediante @Async y un pool de 4 a 8 hilos con capacidad de 500 tareas. Esta cola es volátil: mejora la "
    "latencia del usuario, pero no entrega reintentos durables ni confirmaciones propias de RabbitMQ o Kafka."
)
add_heading("Autenticación multifactor", 2)
add_body(
    "OTP verifica posesión temporal del correo mediante un código de seis dígitos y vigencia limitada. TOTP deriva "
    "el valor a partir de un secreto compartido y el tiempo, con pasos de 30 segundos según RFC 6238. El sistema "
    "genera el secreto, presenta el QR, confirma la vinculación y solicita el código durante el inicio de sesión."
)
add_heading("Almacenamiento híbrido e integridad", 2)
add_body(
    "Separar el binario de los datos transaccionales reduce el crecimiento de PostgreSQL. PocketBase conserva el "
    "archivo, mientras PostgreSQL mantiene el identificador, MIME, tamaño, reporte y SHA-256. El hash no cifra la "
    "imagen: permite detectar cambios y demostrar correspondencia entre registro y contenido."
)
add_heading("Inteligencia artificial local", 2)
add_body(
    "Ollama ejecuta llama3.2 dentro del entorno local, evitando enviar el contexto operacional a una API externa. "
    "El backend construye un contexto permitido, clasifica la intención, aplica límites y controla timeouts. Las "
    "respuestas son asistencia para el administrador y no sustituyen las decisiones de despacho."
)

add_heading("Desarrollo e implementación del sistema", 1)
add_heading("Clúster y ejecución distribuida", 2)
add_figure(
    "estado-swarm.png",
    "Estado verificado del stack Docker Swarm.",
    "La Figura 2 resume la comprobación realizada sobre el stack: dos réplicas de frontend, dos de backend y una "
    "instancia para cada servicio con estado. También se verificaron los endpoints de salud sin exponer secretos.",
    6.35
)
add_figure(
    "docker-swarm-real.png",
    "Contenedores y tareas reales administrados por Docker Swarm.",
    "La Figura 3 corresponde a la captura proporcionada durante la ejecución. Los indicadores verdes muestran las "
    "tareas activas del frontend y backend, mientras las filas detenidas conservan el historial de tareas reemplazadas. "
    "Esta vista permite explicar que Swarm crea una tarea nueva cuando una réplica termina.",
    6.3
)
add_table(
    ["Servicio", "Puerto del host", "Puerto interno", "Acceso"],
    [
        ["Frontend", "8080", "80", "http://localhost:8080"],
        ["Backend", "8082", "8081", "http://localhost:8082/api/health"],
        ["PostgreSQL", "5433", "5432", "Solo administración local."],
        ["PocketBase", "8091", "8090", "http://localhost:8091/_/"],
        ["Ollama", "11435", "11434", "API local del modelo."],
        ["ngrok inspector", "4041", "4040", "http://localhost:4041"],
    ],
    [1.5, 1.2, 1.2, 2.5]
)
add_algorithm(
    1,
    "Despliegue reproducible del stack.",
    """# Preparar variables sin publicar secretos
.\\scripts\\configurar-docker.ps1

# Construir imágenes y desplegar
.\\scripts\\iniciar-swarm.ps1

# Verificar estado deseado
.\\scripts\\estado-swarm.ps1
docker service ls

# Prueba de autorrecuperación
docker service ps gestion-bomberil_backend
docker kill <contenedor_de_una_tarea>
docker service ps gestion-bomberil_backend""",
    "El Algoritmo 1 presenta la secuencia operativa utilizada. Swarm no reinicia el mismo contenedor: reemplaza la "
    "tarea terminada por otra hasta recuperar el número de réplicas definido."
)

add_heading("Interfaz y experiencia de usuario", 2)
add_figure(
    "login.png",
    "Inicio de sesión unificado de la plataforma.",
    "La Figura 4 evidencia la pantalla de acceso. El diseño permite autenticación con Google o credenciales de "
    "operador y mantiene una identidad visual común en modo claro y oscuro.",
    5.9
)
add_figure(
    "ciudadano-reporte.png",
    "Formulario ciudadano de reporte georreferenciado.",
    "La Figura 5 muestra el flujo público para seleccionar el tipo de emergencia, describirla y confirmar la "
    "ubicación. Las categorías visuales reducen errores de clasificación y la validación impide continuar con "
    "datos incompletos.",
    4.2
)
add_figure(
    "operador-espera.png",
    "Unidad operativa en espera de despacho.",
    "La Figura 6 presenta la vista del operador después de seleccionar una unidad. El estado de disponibilidad se "
    "mantiene vinculado a una sesión activa y el canal de despacho informa si la comunicación está disponible.",
    6.25
)
add_figure(
    "operador-asignacion-iot.png",
    "Atención de emergencia, mapa e indicadores IoT.",
    "La Figura 7 evidencia la interfaz durante una asignación: mapa, información del incidente, estado operativo "
    "y métricas de gas. El operador puede finalizar y liberar la unidad cuando completa la bitácora.",
    6.3
)

add_heading("Flujo operativo y consistencia", 2)
add_figure(
    "flujo-operativo.png",
    "Flujo completo desde el reporte hasta el cierre.",
    "La Figura 8 integra los pasos funcionales del sistema. Cada transición persiste primero el estado válido y "
    "después publica la notificación; así, una reconexión puede recuperar la verdad desde PostgreSQL.",
    6.35
)
add_algorithm(
    2,
    "Despacho transaccional de una unidad.",
    """@Transactional
despachar(reporteId, unidadId):
    reporte = bloquearReporte(reporteId)
    unidad  = bloquearUnidad(unidadId)
    validar(reporte.estado == PENDIENTE)
    validar(unidad.estado == DISPONIBLE)
    validar(unidad.operadorActivo)
    reporte.estado = EN_ATENCION
    unidad.estado = EN_RUTA
    guardar(reporte, unidad)
    publicarEventoWebSocket(reporte, unidad)
    return asignacion""",
    "El Algoritmo 2 abstrae el mecanismo de exclusión aplicado. Los bloqueos pesimistas y la transacción evitan que "
    "dos administradores asignen simultáneamente la misma unidad o vuelvan a despachar un reporte atendido."
)
add_table(
    ["Entidad", "Estados principales", "Regla crítica"],
    [
        ["Reporte", "PENDIENTE → EN_ATENCION → ATENDIDO", "ATENDIDO no puede volver a despacharse."],
        ["Unidad", "DISPONIBLE → EN_RUTA → EN_SITIO → DISPONIBLE", "Solo se ofrece si existe operador activo."],
        ["Evaluación IoT", "ACTIVA → FINALIZADA", "Un nodo no se vincula a dos evaluaciones activas."],
        ["Orden Premium", "PENDIENTE → APROBADA/PAGADA o FALLIDA", "Monto y moneda se validan en servidor."],
    ],
    [1.2, 2.4, 2.8]
)

add_heading("Módulo Premium y PayPal Sandbox", 2)
add_figure(
    "premium-formulario.png",
    "Captura de datos para la instalación preventiva.",
    "La Figura 9 muestra el registro del establecimiento previo al pago. Los datos permiten coordinar una visita "
    "técnica y no se consideran pagados hasta recibir y validar la captura de PayPal.",
    5.0
)
add_figure(
    "flujo-paypal.png",
    "Flujo de creación y captura del pago Sandbox.",
    "La Figura 10 detalla el proceso del plan Premium. La orden interna se vincula con la orden de PayPal, y la "
    "confirmación final se almacena solo después de conciliar la respuesta del proveedor.",
    6.35
)
add_algorithm(
    3,
    "Confirmación idempotente de pago.",
    """confirmarPago(ordenInterna, paypalOrderId):
    validar(ordenInterna.paypalOrderId == paypalOrderId)
    if ordenInterna.estado == PAGADO:
        return ordenInterna
    respuesta = capturarEnPayPal(paypalOrderId)
    validar(respuesta.moneda == "USD")
    validar(respuesta.monto == precioServidor)
    validar(respuesta.customId == ordenInterna.codigo)
    if respuesta.estado == "COMPLETED":
        ordenInterna.estado = PAGADO
        ordenInterna.instalacionMaxima = hoy + 2 dias
    guardarAuditoriaOperativa(ordenInterna)
    return ordenInterna""",
    "El Algoritmo 3 resume la defensa contra manipulación de precio y reintentos. La implementación observada "
    "registra la orden y el resultado, aunque el proyecto todavía requiere una vista administrativa específica "
    "con filtros, historial de cambios y reembolsos."
)

add_heading("Correo OTP y TOTP", 2)
add_body(
    "El correo OTP se compone mediante una plantilla HTML y se envía en segundo plano. El código tiene seis dígitos "
    "y expira después de diez minutos. Para TOTP, el backend genera un secreto Base32, construye el QR, confirma la "
    "primera clave y valida ventanas de 30 segundos con tolerancia controlada."
)
add_algorithm(
    4,
    "Validación TOTP durante el inicio de sesión.",
    """validarTotp(usuario, codigo):
    validar(usuario.mfaHabilitado)
    tiempoActual = epochSeconds() / 30
    for desplazamiento in [-1, 0, 1]:
        esperado = HOTP(usuario.secreto, tiempoActual + desplazamiento)
        if comparacionConstante(esperado, codigo):
            emitirJWT(usuario)
            return AUTORIZADO
    registrarIntentoFallido(usuario)
    return CODIGO_INVALIDO_O_VENCIDO""",
    "El Algoritmo 4 representa la validación del segundo factor. La ventana ±1 reduce falsos rechazos por una "
    "pequeña diferencia de reloj. La desactivación autoservicio de TOTP no fue identificada y queda como mejora."
)
add_table(
    ["Control", "Estado", "Evidencia técnica"],
    [
        ["OTP por correo", "Implementado", "Código de 6 dígitos, plantilla HTML, expiración y envío @Async."],
        ["Recuperación de contraseña", "Pendiente", "No se identificó enlace firmado de recuperación."],
        ["Adjuntos de correo", "Pendiente", "No se identificó MimeMessage con archivos adjuntos."],
        ["TOTP activar/validar", "Implementado", "Secreto, QR, confirmación y validación en login."],
        ["TOTP desactivar", "Pendiente", "Debe añadirse endpoint protegido y reautenticación."],
    ],
    [1.8, 1.2, 3.4]
)

add_heading("Almacenamiento privado con PocketBase", 2)
add_figure(
    "flujo-almacenamiento.png",
    "Almacenamiento híbrido de evidencia multimedia.",
    "La Figura 11 explica la separación del binario y sus metadatos. El navegador consume el endpoint autorizado de "
    "Spring Boot y no utiliza directamente la credencial técnica de PocketBase.",
    6.35
)
add_figure(
    "pocketbase-evidencias.png",
    "Registro real de evidencia en PocketBase.",
    "La Figura 12 evidencia una imagen almacenada con SHA-256, identificador de reporte, MIME y tamaño. El token "
    "mostrado durante la configuración nunca se incorpora al documento ni al repositorio.",
    6.2
)
add_algorithm(
    5,
    "Carga e integridad de la evidencia.",
    """guardarEvidencia(reporteId, archivo):
    validarMimeYTamano(archivo)
    hash = SHA256(archivo.bytes)
    registroPB = pocketBase.crear(
        archivo, reporteId, hash, mime, tamano)
    evidenciaSQL = nueva Evidencia(
        reporteId, registroPB.id, hash, mime, tamano)
    repositorio.guardar(evidenciaSQL)
    return evidenciaSQL

consultarEvidencia(id, usuario):
    validarRolYAcceso(usuario, id)
    return backend.descargarDesdePocketBase(id)""",
    "El Algoritmo 5 presenta la escritura y lectura controlada. Las reglas bloqueadas evitan una consulta anónima "
    "del registro; como endurecimiento adicional se recomienda marcar el campo como archivo protegido y usar tokens "
    "de corta duración o mantener exclusivamente el proxy actual."
)

add_heading("Integración IoT ESP32 post-incendio", 2)
add_body(
    "El nodo táctico se diseñó para colocarse después de mitigar un incendio estructural. El sensor MQ-2 entrega "
    "una lectura analógica; el firmware clasifica RESPIRABLE, PRECAUCIÓN o CRÍTICO, actualiza LCD y LED RGB, y activa "
    "el buzzer cuando se supera el umbral peligroso. En paralelo publica telemetría con nodo, valor y evento."
)
add_figure(
    "iot-wokwi.png",
    "Prototipo virtual del nodo táctico en Wokwi.",
    "La Figura 13 muestra la simulación del circuito: ESP32, sensor MQ-2, pantalla I2C, LED RGB, pulsador y buzzer. "
    "La simulación permitió verificar pines, estados visuales y umbrales antes de energizar el montaje real.",
    6.35
)
add_figure(
    "iot-prototipo-fisico.png",
    "Implementación física del nodo IoT post-incendio.",
    "La Figura 14 presenta el prototipo físico funcional con ESP32, protoboard, sensor, indicadores y LCD integrados "
    "en una carcasa de demostración. Esta evidencia confirma el paso de la simulación al dispositivo real.",
    6.1
)
add_figure(
    "iot-dashboard.png",
    "Monitoreo de telemetría IoT en la interfaz del operador.",
    "La Figura 15 muestra la recepción de gas y el estado de transmisión. Para evitar estados falsos, la interfaz "
    "considera conectado al nodo únicamente cuando existe telemetría reciente; si expira el intervalo, informa "
    "desconexión aunque las lecturas históricas permanezcan visibles.",
    6.2
)
add_algorithm(
    6,
    "Publicación periódica de telemetría desde ESP32.",
    """loop():
    gas = analogRead(PIN_GAS)
    estado = clasificar(gas, UMBRAL_ADVERTENCIA, UMBRAL_PELIGRO)
    actualizarLCD_RGB_Buzzer(estado)
    if sistemaEncendido and millis() - ultimoEnvio >= intervalo:
        POST /api/iot/telemetria
        headers: X-IOT-KEY
        body: { nodoId, nivelGas: gas, evento: "TELEMETRIA" }
        registrarCodigoHTTP()
    reconectarWiFiSiCorresponde()""",
    "El Algoritmo 6 evita bloquear el ciclo con una alarma basada en millis(). Los umbrales residen en el firmware "
    "y el mismo estado calculado se envía al backend, lo que impide que dispositivo e interfaz interpreten valores "
    "con reglas diferentes."
)

add_heading("Integración de inteligencia artificial local", 2)
add_body(
    "El Copiloto Operativo consulta únicamente datos del sistema autorizados: emergencias activas, unidades, "
    "operadores, despachos, bitácoras y monitoreo IoT. Las preguntas fuera de dominio se rechazan. Para evitar una "
    "respuesta puramente determinista, el backend recupera contexto validado y construye un prompt acotado que "
    "Ollama transforma en una explicación legible."
)
add_table(
    ["Puede responder", "No debe responder"],
    [
        ["Unidades disponibles y su tipo.", "Programación general o tareas académicas ajenas."],
        ["Resumen del turno y emergencias activas.", "Datos personales no necesarios para la operación."],
        ["Estado reciente del nodo IoT y habitabilidad.", "Instrucciones que alteren seguridad o credenciales."],
        ["Estado de despachos y cierres operativos.", "Información no contenida en el contexto del sistema."],
    ],
    [3.2, 3.2]
)
add_algorithm(
    7,
    "Consulta acotada del Copiloto Operativo.",
    """responder(pregunta, administrador):
    validarJWTyRol(administrador)
    intencion = clasificarIntencion(pregunta)
    if intencion not in DOMINIO_OPERATIVO:
        return respuestaFueraDeAlcance()
    contexto = consultarDatosMinimos(intencion)
    prompt = construirPromptSistema(contexto, pregunta)
    respuesta = ollama.chat("llama3.2", prompt, timeout)
    validarRespuesta(respuesta)
    return sanitizar(respuesta)""",
    "El Algoritmo 7 reduce alucinaciones al consultar primero el estado real y limitar el modelo. Si Ollama excede "
    "el tiempo de espera o devuelve una respuesta inválida, la API entrega un error controlado y no bloquea el dashboard."
)

add_heading("Integración con contenidos anteriores", 1)
add_table(
    ["Contenido conservado", "Aplicación actual"],
    [
        ["Cliente–servidor", "Angular consume una API Spring Boot desacoplada."],
        ["WebSocket", "STOMP/SockJS publica reportes, despachos, cierres e IoT."],
        ["Concurrencia", "Pool asíncrono y bloqueos pesimistas en asignación."],
        ["Excepciones", "Manejador global devuelve errores JSON consistentes."],
        ["Logs", "Eventos de autenticación, despacho, IoT, correo y pagos."],
        ["JWT / OAuth 2.0", "Rutas protegidas y autenticación Google."],
        ["Roles", "Ciudadano público; operador y administrador autenticados."],
        ["SOLID", "Controladores delgados, servicios cohesionados y repositorios."],
        ["Patrones", "Repository, Service Layer, DTO, Strategy de almacenamiento y Observer pub/sub."],
        ["Transacciones", "Estados operativos, asignación, cierre y pago."],
    ],
    [1.65, 4.75]
)
add_body(
    "El diseño sigue el principio de responsabilidad única: los controladores traducen HTTP, los servicios aplican "
    "reglas y los repositorios encapsulan persistencia. EvidenceStorageService actúa como abstracción de almacenamiento "
    "y permite seleccionar PocketBase o almacenamiento local mediante configuración. Las integraciones externas se "
    "encapsulan para que sus errores no contaminen directamente la lógica de presentación."
)

add_heading("Pruebas realizadas", 1)
add_table(
    ["Nivel", "Prueba", "Resultado"],
    [
        ["Backend", "Maven test: controladores, contexto, Ollama y servicios de unidades.", "12/12 aprobadas."],
        ["Frontend", "Angular/Karma: componentes y servicios.", "7/7 aprobadas."],
        ["Clúster", "Estado de servicios y réplicas.", "2/2 frontend, 2/2 backend; demás 1/1."],
        ["Salud", "Frontend, backend, PocketBase, Ollama y ngrok.", "HTTP 200 / backend UP."],
        ["Persistencia", "Conteos y estados en PostgreSQL.", "Datos consistentes y consultables."],
        ["IoT", "WiFi, DNS, HTTPS, autenticación técnica y telemetría.", "Flujo funcional con clave coincidente."],
    ],
    [1.25, 3.45, 1.7]
)
add_body(
    "La base consultada durante la verificación contenía 6 reportes ciudadanos, 6 unidades, 2 evidencias multimedia, "
    "1 orden Premium y 85 lecturas IoT. Cuatro reportes estaban ATENDIDOS y dos PENDIENTES; la orden registrada figuraba "
    "PAGADA. Estas cifras son una instantánea de prueba y pueden variar con nuevas ejecuciones."
)
add_status_box("APROBADO", "Backend: 12 pruebas, sin fallos, errores ni omisiones.", GREEN)
add_status_box("APROBADO", "Frontend: 7 pruebas distribuidas en 4 archivos, todas exitosas.", GREEN)
add_status_box("VERIFICADO", "Servicios del stack operativos y endpoints de salud accesibles.", BLUE)

add_heading("Resultados obtenidos", 1)
add_bullets([
    "El administrador recibe y gestiona reportes; el operador obtiene la asignación sin recargar gracias a eventos y reconciliación.",
    "Solo se ofrecen para despacho unidades con operador activo y estado DISPONIBLE.",
    "El cierre publica el resultado en tiempo real, libera la unidad y bloquea un nuevo despacho del reporte ATENDIDO.",
    "Las evidencias se almacenan fuera de PostgreSQL y conservan hash SHA-256 y referencia relacional.",
    "El pago Sandbox genera una orden Premium y programa la instalación después de la captura válida.",
    "La telemetría real del ESP32 se incorpora al reporte con primera y última toma de aire.",
    "Ollama responde dentro del dominio operacional y utiliza el estado del sistema como contexto."
])

add_heading("Problemas encontrados y soluciones aplicadas", 1)
add_table(
    ["Problema", "Causa", "Solución aplicada / recomendada"],
    [
        ["Asignación no aparecía hasta recargar.", "Eventos entre réplicas con broker simple.", "WebSocket más reconciliación REST y eventos de cierre."],
        ["IoT mostraba conectado sin transmitir.", "Se usaba la existencia de lecturas históricas.", "Estado por antigüedad de la última telemetría."],
        ["HTTP 401 del ESP32.", "Clave técnica distinta entre firmware y stack.", "Unificar IOT_NODE_KEY y reiniciar el servicio."],
        ["HTTPS 421 / conexión rechazada.", "Ruta/túnel/red y SNI.", "Diagnóstico DNS/TLS, URL HTTPS y túnel correcto."],
        ["OAuth origin_mismatch.", "Origen localhost no autorizado.", "Registrar el origen exacto en Google Cloud."],
        ["Multimedia ausente.", "Configuración/token PocketBase y referencias.", "Proveedor PocketBase, colección, proxy y SHA-256."],
        ["Réplicas 0/2 temporalmente.", "Base de datos no saludable o recursos.", "Healthchecks, orden de arranque y revisión de tareas."],
    ],
    [1.7, 2.0, 2.8]
)

add_heading("Matriz de cumplimiento del enunciado", 1)
add_table(
    ["Requisito", "Estado", "Observación verificable"],
    [
        ["Docker Swarm y réplicas", "Cumple", "Frontend 2, backend 2 y PostgreSQL 1."],
        ["REST síncrono", "Cumple", "46 endpoints mapeados aproximadamente."],
        ["Cola de mensajes externa", "Parcial", "Existe cola en memoria; falta RabbitMQ/Kafka durable."],
        ["Pagos Sandbox", "Cumple parcial", "Pago y conciliación listos; falta consola admin completa/auditoría de cambios."],
        ["OTP por correo", "Cumple", "HTML, expiración y ejecución @Async."],
        ["Recuperación y adjuntos", "Pendiente", "No fueron identificados en el código revisado."],
        ["TOTP", "Cumple parcial", "Alta y validación listas; falta desactivación autoservicio."],
        ["PocketBase privado", "Cumple", "Acceso vía backend, reglas bloqueadas, SHA-256."],
        ["Inteligencia artificial", "Cumple", "Ollama llama3.2 con dominio acotado."],
        ["Soft delete, paginación y filtros", "Parcial", "Estados conservan operaciones; faltan patrones uniformes en todos los listados."],
        ["README y .env.example", "Cumple", "Guías Docker y ejemplo sin secretos."],
    ],
    [2.0, 1.1, 3.45]
)
add_body(
    "La matriz evita presentar como terminadas funciones que no aparecen en la implementación. Para alcanzar el 100 % "
    "del enunciado se debe incorporar un broker durable, recuperación de contraseña, adjuntos de correo, desactivación "
    "TOTP, consola administrativa de transacciones con auditoría y una estrategia uniforme de borrado lógico/paginación."
)

add_heading("Conclusiones finales", 1)
add_body(
    "El proyecto integra sus componentes en un ecosistema distribuido funcional. La plataforma "
    "coordina actores, datos relacionales, archivos, pagos, IA e IoT sin perder el ciclo operacional del incidente. "
    "El uso de réplicas demuestra autorrecuperación frente a fallos de tareas y el empleo de una sola base conserva "
    "una fuente de verdad coherente."
)
add_body(
    "La integración ESP32 constituye el aporte más tangible: las proyecciones del informe anterior se materializaron "
    "en un dispositivo físico que mide el ambiente post-incendio y entrega sus datos al informe del operador. La "
    "comparación Wokwi–prototipo confirma una metodología incremental de simulación, construcción y validación."
)
add_body(
    "La revisión también evidencia que distribuir no significa únicamente añadir contenedores. WebSocket con broker "
    "local, una cola volátil y un clúster de un solo nodo tienen límites claros. Reconocerlos permite proponer una "
    "evolución responsable hacia RabbitMQ, relay STOMP, múltiples nodos físicos, observabilidad y respaldos coordinados."
)

add_heading("Recomendaciones", 1)
add_bullets([
    "Incorporar RabbitMQ como cola durable y como relay STOMP compartido entre réplicas.",
    "Añadir un segundo nodo físico al Swarm para tolerar la pérdida del anfitrión.",
    "Completar recuperación de contraseña, adjuntos HTML y desactivación TOTP protegida.",
    "Crear una consola administrativa de pagos con filtros, auditoría, conciliación y reembolsos.",
    "Uniformar borrado lógico, restauración, paginación y filtros en usuarios, pagos y archivos.",
    "Agregar métricas Prometheus/Grafana, trazas y alertas de salud.",
    "Calibrar el MQ-2 con condiciones controladas; no usar el prototipo como certificación profesional de habitabilidad."
])

add_heading("Referencias bibliográficas", 1)
refs = [
    "Docker, Inc. (2026). How services work. https://docs.docker.com/engine/swarm/how-swarm-mode-works/services/",
    "Docker, Inc. (2026). Use Swarm mode routing mesh. https://docs.docker.com/engine/swarm/ingress/",
    "Spring. (2026). WebSocket STOMP overview. https://docs.spring.io/spring-framework/reference/web/websocket/stomp/overview.html",
    "Spring. (2026). Task execution and scheduling. https://docs.spring.io/spring-framework/reference/integration/scheduling.html",
    "PayPal. (2026). Orders API v2. https://developer.paypal.com/docs/api/orders/v2/",
    "PocketBase. (2026). Files upload and handling. https://pocketbase.io/docs/files-handling/",
    "PocketBase. (2026). API rules and filters. https://pocketbase.io/docs/api-rules-and-filters/",
    "M'Raihi, D., Machani, S., Pei, M., & Rydell, J. (2011). TOTP: Time-Based One-Time Password Algorithm (RFC 6238). https://datatracker.ietf.org/doc/html/rfc6238",
    "Fette, I., & Melnikov, A. (2011). The WebSocket Protocol (RFC 6455). https://datatracker.ietf.org/doc/html/rfc6455",
    "Ollama. (2026). API documentation. https://docs.ollama.com/api",
    "Angular. (2026). Angular documentation. https://angular.dev/",
    "PostgreSQL Global Development Group. (2026). PostgreSQL documentation. https://www.postgresql.org/docs/",
]
for ref in refs:
    p = doc.add_paragraph(style="Bibliography" if "Bibliography" in [s.name for s in doc.styles] else "Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(6)
    p.add_run(ref)

add_heading("Anexos", 1)
add_heading("Anexo A. Repositorio y archivos de ejecución", 2)
add_body(
    "Repositorio: https://github.com/Joseph-Franco692/Proyecto_Emergencias_jf. Los archivos docker-stack.yml, "
    "compose.yaml, DOCKER_SWARM.md, POCKETBASE_EVIDENCIAS.md, .env.docker.example y los scripts de la carpeta "
    "scripts documentan el despliegue. El archivo .env.docker real se mantiene excluido de Git."
)
add_heading("Anexo B. Comandos de verificación", 2)
add_algorithm(
    8,
    "Verificación integral sin exponer secretos.",
    """# Servicios y tareas
docker service ls
docker stack services gestion-bomberil
docker stack ps gestion-bomberil

# Salud
Invoke-WebRequest http://localhost:8082/api/health
Invoke-WebRequest http://localhost:8091/api/health
Invoke-WebRequest http://localhost:11435/api/tags

# Pruebas
cd emergencias
mvn test
cd ..\\central-bomberos
npm test -- --watch=false""",
    "El Algoritmo 8 reúne la evidencia reproducible de clúster, salud y pruebas. Ningún comando imprime variables "
    "sensibles ni el contenido del archivo de entorno."
)
add_heading("Anexo C. Correspondencia de pines del prototipo IoT", 2)
add_table(
    ["Elemento", "Pin ESP32", "Función"],
    [
        ["MQ-2 AOUT", "GPIO 34", "Entrada analógica de humo/gas."],
        ["Buzzer", "GPIO 23", "Alarma acústica."],
        ["Pulsador", "GPIO 4", "Encendido/apagado lógico."],
        ["RGB rojo", "GPIO 5", "Estado crítico."],
        ["RGB verde", "GPIO 19", "Estado respirable."],
        ["RGB azul", "GPIO 18", "Sistema en espera/apagado."],
        ["LCD SDA", "GPIO 21", "Bus I2C de datos."],
        ["LCD SCL", "GPIO 22", "Bus I2C de reloj."],
    ],
    [2.1, 1.25, 3.2]
)
add_body(
    "Los valores de umbral son parámetros de calibración del prototipo y deben ajustarse con mediciones de referencia. "
    "El dispositivo es una demostración académica; no reemplaza instrumentos certificados ni protocolos profesionales "
    "para declarar habitable una estructura."
)

# Word should update TOC, table of figures and sequence fields when opened.
settings = doc.settings._element
update = settings.find(qn("w:updateFields"))
if update is None:
    update = OxmlElement("w:updateFields")
    settings.append(update)
update.set(qn("w:val"), "true")

core = doc.core_properties
core.title = "Sistema distribuido de gestión bomberil — Informe completo"
core.subject = "Aplicaciones Distribuidas — informe técnico unificado"
core.author = "Joseph Franco"
core.keywords = "Docker Swarm, Spring Boot, Angular, WebSocket, IoT, PocketBase, PayPal, Ollama"

doc.save(OUT)
print(OUT)
