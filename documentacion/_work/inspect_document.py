from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def text_of_cell(cell) -> str:
    return " | ".join(
        paragraph.text.strip()
        for paragraph in cell.paragraphs
        if paragraph.text.strip()
    )


def main() -> None:
    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    doc = Document(source)

    paragraphs = []
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:
            paragraphs.append(
                {
                    "index": index,
                    "style": paragraph.style.name if paragraph.style else None,
                    "text": text,
                }
            )

    tables = []
    for table_index, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([text_of_cell(cell) for cell in row.cells])
        tables.append({"index": table_index, "rows": rows})

    sections = []
    for index, section in enumerate(doc.sections):
        sections.append(
            {
                "index": index,
                "page_width": section.page_width,
                "page_height": section.page_height,
                "top_margin": section.top_margin,
                "bottom_margin": section.bottom_margin,
                "left_margin": section.left_margin,
                "right_margin": section.right_margin,
                "different_first_page": section.different_first_page_header_footer,
            }
        )

    styles = {}
    for style in doc.styles:
        if style.type != 1:
            continue
        paragraph_format = style.paragraph_format
        font = style.font
        styles[style.name] = {
            "font_name": font.name,
            "font_size": font.size.pt if font.size else None,
            "bold": font.bold,
            "italic": font.italic,
            "alignment": paragraph_format.alignment,
            "space_before": paragraph_format.space_before.pt
            if paragraph_format.space_before
            else None,
            "space_after": paragraph_format.space_after.pt
            if paragraph_format.space_after
            else None,
            "line_spacing": paragraph_format.line_spacing,
        }

    body = doc.element.body
    body_order = []
    paragraph_index = 0
    table_index = 0
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            body_order.append({"kind": "paragraph", "index": paragraph_index})
            paragraph_index += 1
        elif child.tag == qn("w:tbl"):
            body_order.append({"kind": "table", "index": table_index})
            table_index += 1

    output.write_text(
        json.dumps(
            {
                "source": str(source),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "inline_shape_count": len(doc.inline_shapes),
                "paragraphs": paragraphs,
                "tables": tables,
                "sections": sections,
                "styles": styles,
                "body_order": body_order,
            },
            ensure_ascii=False,
            indent=2,
            default=str,
        ),
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
